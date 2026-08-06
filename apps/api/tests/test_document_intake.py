"""Document intake + "document" provenance (v6.0 Phase 6).

Pins:
  1. extraction works PER FORMAT against real generated files — .docx
     (paragraphs AND table cells), .txt/.md, and PDF — with honest,
     machine-coded refusals for renamed files, unsupported types,
     text-less/image-only PDFs, and binary-in-.txt;
  2. a number that appears in an ingested document PASSES the invoice
     line gate stamped "document" and the proposal price/number gates;
  3. a number NOT in the document still REFUSES — ingestion widens the
     sanctioned set, it never disables the gate;
  4. reading a document grounds the run (source.md) and records the
     file's sha256 + numbers for the receipt.
"""
import asyncio
import json

import pytest

from app.services import document_service, operator_service, state_store
from app.services import operator_tools as t
from app.services.operator_context import OperatorContext, set_current_operator

_CUSTOMERS = [{"id": "42", "name": "Sandy Alvarez", "email": "sandy@gulf.test"}]

# The client's RFP text. 7200 is the budget line; 1350 lives only in a table.
_RFP_TEXT = (
    "Gulf Realty — Request for Proposal\n\n"
    "We are seeking an AI automation partner for intake and scheduling.\n"
    "Our approved budget for this engagement is $7,200.\n"
    "Work should begin within 3 weeks of signature.\n")
_TABLE_ROWS = [("Item", "Rate"), ("Monthly retainer", "$1,350")]


@pytest.fixture(autouse=True)
def _isolated_state(monkeypatch, tmp_path):
    monkeypatch.setattr(state_store, "STATE_DIR", tmp_path / "state")


def _op(tmp_path, stated=None):
    async def _emit(_ev):
        return None
    record = {"id": "op_doc", "steps": [], "tools_used": [], "artifacts": [],
              "errors": [], "user_stated_numbers": list(stated or [])}
    op = OperatorContext(folder=tmp_path, record=record, emit=_emit)
    set_current_operator(op)
    return op


def _call(_tool_name, **kwargs):
    tool = next(x for x in t.PLANNER_TOOLS if x.name == _tool_name)
    raw = asyncio.run(tool.call(kwargs))
    return json.loads(raw) if isinstance(raw, str) else raw


@pytest.fixture()
def qb(monkeypatch):
    created = []
    monkeypatch.setattr(t.quickbooks_service, "list_customers", lambda: list(_CUSTOMERS))
    monkeypatch.setattr(t.quickbooks_service, "list_items", lambda: [])

    def fake_create(customer_id, lines, txn_date="", due_date=""):
        created.append({"customer_id": customer_id, "lines": lines})
        return {"id": "99", "doc_number": "1042", "customer": "Sandy Alvarez",
                "total": 7200.0, "email_status": "NotSet", "link": ""}

    monkeypatch.setattr(t.quickbooks_service, "create_invoice", fake_create)
    return created


def _make_docx(path, text=_RFP_TEXT, rows=_TABLE_ROWS):
    import docx
    d = docx.Document()
    for line in text.splitlines():
        d.add_paragraph(line)
    if rows:
        table = d.add_table(rows=len(rows), cols=len(rows[0]))
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                table.cell(r, c).text = val
    d.save(str(path))
    return path


def _make_pdf(path, text=_RFP_TEXT):
    """A minimal single-page PDF with a real text layer (no libs beyond
    what's already vendored — pypdf reads what we hand-assemble)."""
    # A base-14 Helvetica content stream is latin-1; the fixture text is
    # transliterated rather than the builder pretending to handle Unicode.
    text = text.replace("—", "-").encode("latin-1", "replace").decode("latin-1")
    lines = [ln for ln in text.splitlines() if ln.strip()]
    content = "BT /F1 12 Tf 40 750 Td 14 TL\n" + "".join(
        f"({ln.replace('(', '').replace(')', '')}) Tj T*\n" for ln in lines) + "ET"
    objs = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        "/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        f"<< /Length {len(content)} >>\nstream\n{content}\nendstream",
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, body in enumerate(objs, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n{body}\nendobj\n".encode("latin-1")
    xref_at = len(out)
    out += f"xref\n0 {len(objs) + 1}\n0000000000 65535 f \n".encode("latin-1")
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode("latin-1")
    out += (f"trailer\n<< /Size {len(objs) + 1} /Root 1 0 R >>\nstartxref\n"
            f"{xref_at}\n%%EOF\n").encode("latin-1")
    path.write_bytes(bytes(out))
    return path


# --------------------------------------------------------------------------
# 1. Extraction per format + honest refusals
# --------------------------------------------------------------------------

def test_docx_extraction_includes_paragraphs_and_table_cells(tmp_path):
    doc = document_service.extract_path(str(_make_docx(tmp_path / "rfp.docx")))
    assert doc["kind"] == "docx" and doc["name"] == "rfp.docx"
    assert "approved budget for this engagement is $7,200" in doc["text"]
    assert "$1,350" in doc["text"]                   # the TABLE rate survives
    assert len(doc["sha256"]) == 64


def test_txt_and_md_extraction(tmp_path):
    p = tmp_path / "rfp.txt"
    p.write_text(_RFP_TEXT, encoding="utf-8")
    assert document_service.extract_path(str(p))["kind"] == "txt"
    m = tmp_path / "notes.md"
    m.write_text("# Scope\n\n" + _RFP_TEXT, encoding="utf-8")
    out = document_service.extract_path(str(m))
    assert out["kind"] == "txt" and "$7,200" in out["text"]


def test_pdf_extraction(tmp_path):
    doc = document_service.extract_path(str(_make_pdf(tmp_path / "rfp.pdf")))
    assert doc["kind"] == "pdf" and doc["pages"] == 1
    assert "7,200" in doc["text"]


@pytest.mark.parametrize("name,data,reason", [
    ("thing.rtf", b"x" * 200, "unsupported"),
    ("empty.txt", b"", "empty"),
    ("fake.docx", b"%PDF-1.4 not really a docx" + b"x" * 100, "not_docx"),
    ("fake.pdf", b"PK\x03\x04 actually a zip" + b"x" * 100, "not_pdf"),
    ("binary.txt", b"text\x00\x00binary" + b"\x00" * 100, "not_text"),
    ("thin.txt", b"too short", "no_text"),
])
def test_refusals_are_honest_and_machine_coded(name, data, reason):
    with pytest.raises(document_service.DocumentError) as exc:
        document_service.extract_bytes(data, name)
    assert exc.value.reason == reason


def test_image_only_pdf_refuses_with_no_text(tmp_path):
    """A real PDF with no text layer is refused, never guessed at."""
    p = _make_pdf(tmp_path / "scan.pdf", text="")
    with pytest.raises(document_service.DocumentError) as exc:
        document_service.extract_path(str(p))
    assert exc.value.reason == "no_text"


def test_missing_file_refuses(tmp_path):
    with pytest.raises(document_service.DocumentError) as exc:
        document_service.extract_path(str(tmp_path / "nope.pdf"))
    assert exc.value.reason == "missing"


# --------------------------------------------------------------------------
# 2 + 4. read_document grounds the run and records provenance
# --------------------------------------------------------------------------

def test_read_document_grounds_the_run_and_records_numbers(tmp_path):
    op = _op(tmp_path)
    out = _call("read_document", path=str(_make_docx(tmp_path / "rfp.docx")))
    assert out["kind"] == "docx" and out["numbers_found"] >= 3
    rec = op.record
    assert 7200.0 in rec["document_numbers"] and 1350.0 in rec["document_numbers"]
    assert rec["documents"][0]["name"] == "rfp.docx"
    assert len(rec["documents"][0]["sha256"]) == 64
    assert rec["grounding_ok"] is True
    body = (tmp_path / "source.md").read_text(encoding="utf-8")
    assert "Document: rfp.docx" in body and "$7,200" in body


def test_read_document_failure_records_nothing(tmp_path):
    op = _op(tmp_path)
    out = _call("read_document", path=str(tmp_path / "missing.pdf"))
    assert out["reason"] == "missing"
    assert not op.record.get("document_numbers")
    assert not (tmp_path / "source.md").exists()


# --------------------------------------------------------------------------
# 2 + 3. THE provenance rule: document numbers pass, others still refuse
# --------------------------------------------------------------------------

def test_document_number_passes_the_invoice_line_gate(tmp_path, qb):
    """$7,200 from the client's own RFP — never typed by the operator —
    satisfies the line gate, stamped "document"."""
    op = _op(tmp_path)                               # nothing user-stated
    _call("read_document", path=str(_make_docx(tmp_path / "rfp.docx")))
    out = _call("create_quickbooks_invoice", customer="Sandy Alvarez",
                lines=[{"description": "Discovery engagement", "amount": 7200}])
    assert out.get("reason") == "invoice_plan_pending"   # gate satisfied → asks
    operator_service._apply_invoice_answer(op, t.INVOICE_PROCEED)
    done = _call("create_quickbooks_invoice", customer="Sandy Alvarez",
                 lines=[{"description": "Discovery engagement", "amount": 7200}])
    assert done["doc_number"] == "1042"
    assert qb[0]["lines"][0]["amount"] == 7200.0


def test_table_only_number_also_passes(tmp_path, qb):
    """The rate lived only in a DOCX table cell — extraction must have
    reached it for this to work."""
    _op(tmp_path)
    _call("read_document", path=str(_make_docx(tmp_path / "rfp.docx")))
    out = _call("create_quickbooks_invoice", customer="Sandy Alvarez",
                lines=[{"description": "Monthly retainer", "amount": 1350}])
    assert out.get("reason") == "invoice_plan_pending"


def test_number_not_in_the_document_still_refuses(tmp_path, qb):
    """THE required pin: ingestion widens the sanctioned set; it does not
    disable the gate."""
    op = _op(tmp_path)
    _call("read_document", path=str(_make_docx(tmp_path / "rfp.docx")))
    out = _call("create_quickbooks_invoice", customer="Sandy Alvarez",
                lines=[{"description": "Discovery engagement", "amount": 9999}])
    assert out.get("reason") == "line_value_unverified"
    assert "isn't in any document I've read" in op.record["needs_input"][-1]["question"]
    assert qb == []


def test_without_any_document_the_gate_is_unchanged(tmp_path, qb):
    op = _op(tmp_path)                               # no document read
    out = _call("create_quickbooks_invoice", customer="Sandy Alvarez",
                lines=[{"description": "Discovery", "amount": 7200}])
    assert out.get("reason") == "line_value_unverified"
    # No document read → the message doesn't claim to have checked documents.
    assert "isn't in any document" not in op.record["needs_input"][-1]["question"]
    assert qb == []


def test_document_price_passes_the_proposal_gate(tmp_path, monkeypatch):
    op = _op(tmp_path)
    _call("add_contact", name="Sandy Alvarez", email="sandy@gulf.test")
    _call("add_deal", contact="Sandy Alvarez", title="AI discovery", stage="proposal")
    _call("read_document", path=str(_make_docx(tmp_path / "rfp.docx")))

    async def fake_agent(system, user_input, **kw):
        return "# Proposal\n\n## Price\n$7,200.00, fixed.\n"

    monkeypatch.setattr(t, "run_text_agent", fake_agent)
    out = _call("draft_proposal", deal="Sandy", price="7200")
    assert out.get("reason") == "proposal_plan_pending"   # RFP price accepted
    # A price in NO document and NOT stated still refuses.
    bad = _call("draft_proposal", deal="Sandy", price="8888")
    assert bad.get("reason") == "price_unverified"
    assert "isn't in any document I've read" in op.record["needs_input"][-1]["question"]