"""Document intake — extract text from PDF / DOCX / TXT (v6.0 Phase 6).

Same discipline as pdf_service (which this delegates to for PDFs): a
document grounds a run ONLY if it yields real extractable text, the bytes
must actually match the format's magic header (a renamed file never
passes), there is a hard size cap, and failure is HONEST — no OCR, no
guessing. A file with no text layer is refused with a machine reason,
never silently treated as empty.

Extracted text becomes a provenance source stamped "document": numbers
that appear in a client's own RFP become citable by the proposal and
invoice gates, exactly the way "deal-record" was added. Numbers that do
NOT appear in the document still refuse.
"""

from __future__ import annotations

import hashlib
import io
import logging
from pathlib import Path

from .pdf_service import MAX_TEXT_CHARS, MIN_TEXT_CHARS, PdfError
from .pdf_service import validate_and_extract as _extract_pdf

log = logging.getLogger("ridian.documents")

MAX_DOC_BYTES = 25 * 1024 * 1024
SUPPORTED_SUFFIXES = (".pdf", ".docx", ".txt", ".md")

# DOCX is a ZIP container: every real one starts with the PK local-file
# header. This rejects a renamed .doc/.pdf/image the same way the %PDF-
# check does for PDFs.
_ZIP_MAGIC = (b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")


class DocumentError(Exception):
    """Renderer-safe intake failure. ``reason`` is a stable machine code
    (unsupported | missing | empty | too_large | not_docx | no_text |
    unreadable | not_text) — plus every pdf_service reason for PDFs."""

    def __init__(self, detail: str, *, reason: str = "invalid", status: int = 400):
        self.detail = detail
        self.reason = reason
        self.status = status
        super().__init__(detail)


def _finish(text: str, *, kind: str, name: str, data: bytes,
            pages: int = 0) -> dict:
    text = (text or "").strip()
    if len(text) < MIN_TEXT_CHARS:
        raise DocumentError(
            f"That {kind.upper()} has no extractable text — I won't guess at its "
            "contents. Paste the text instead, or provide a text-bearing file.",
            reason="no_text")
    truncated = len(text) > MAX_TEXT_CHARS
    if truncated:
        text = text[:MAX_TEXT_CHARS].rstrip() + "\n\n[... truncated ...]"
    log.info("document.extracted name=%s kind=%s chars=%d", name, kind, len(text))
    return {"text": text, "kind": kind, "name": name, "chars": len(text),
            "pages": pages, "truncated": truncated,
            "sha256": hashlib.sha256(data).hexdigest()}


def _extract_docx(data: bytes, name: str) -> dict:
    head = data[:4]
    if not any(head.startswith(m) for m in _ZIP_MAGIC):
        raise DocumentError(
            "That file isn't a .docx (it's missing the ZIP container header). "
            "Check the file and try again.", reason="not_docx")
    try:
        import docx  # noqa: PLC0415 — real dependency (python-docx)
        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]
        # Tables carry the numbers that matter in an RFP — a rate card in a
        # table must be as citable as one in a paragraph.
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
    except DocumentError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise DocumentError(
            f"Couldn't read that .docx ({type(exc).__name__}) — it may be corrupt "
            "or password-protected.", reason="unreadable") from exc
    text = "\n".join(p.strip() for p in parts if p and p.strip())
    return _finish(text, kind="docx", name=name, data=data)


def _extract_txt(data: bytes, name: str) -> dict:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = data.decode("cp1252")
        except UnicodeDecodeError as exc:
            raise DocumentError(
                "That file isn't readable text (it isn't UTF-8 or Windows-1252). "
                "If it's a binary document, give it its real extension.",
                reason="not_text") from exc
    # A binary blob renamed .txt usually decodes with NUL bytes — refuse it
    # rather than feeding control characters to the model.
    if "\x00" in text:
        raise DocumentError(
            "That .txt contains binary data, not text.", reason="not_text")
    return _finish(text, kind="txt", name=name, data=data)


def extract_bytes(data: bytes, filename: str) -> dict:
    """Extract text from raw bytes, dispatching on the filename's suffix."""
    name = Path(str(filename or "document")).name
    suffix = Path(name).suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise DocumentError(
            f"Unsupported file type {suffix or '(none)'} — I can read "
            f"{', '.join(SUPPORTED_SUFFIXES)}.", reason="unsupported")
    if not data:
        raise DocumentError("That file is empty.", reason="empty")
    if len(data) > MAX_DOC_BYTES:
        raise DocumentError(
            f"That file is too large ({len(data) // (1024 * 1024)} MB; the limit "
            f"is {MAX_DOC_BYTES // (1024 * 1024)} MB).", reason="too_large")
    if suffix == ".pdf":
        try:
            out = _extract_pdf(data, name)
        except PdfError as exc:   # preserve the machine reason across the seam
            raise DocumentError(exc.detail, reason=exc.reason) from exc
        return {**out, "kind": "pdf", "name": name,
                "sha256": hashlib.sha256(data).hexdigest()}
    if suffix == ".docx":
        return _extract_docx(data, name)
    return _extract_txt(data, name)


def extract_path(path: str) -> dict:
    """Extract text from a file on disk. Refuses missing files honestly."""
    p = Path(str(path or "").strip('"').strip())
    if not p.exists() or not p.is_file():
        raise DocumentError(f"No file at {p}.", reason="missing")
    try:
        data = p.read_bytes()
    except OSError as exc:
        raise DocumentError(f"Couldn't read {p.name}: {exc}",
                            reason="unreadable") from exc
    return {**extract_bytes(data, p.name), "path": str(p)}
